import eventlet
eventlet.monkey_patch()

import os
import json
import tempfile
import uuid
import re
import shutil
import traceback
import hashlib
from datetime import datetime
from flask import Flask, render_template, request, jsonify, session, send_file, abort
from flask_socketio import SocketIO, join_room
from flask_cors import CORS
from werkzeug.utils import secure_filename
import google.generativeai as genai
from docx import Document
from dotenv import load_dotenv
import pythoncom
from celery import Celery, Task
from sqlalchemy.orm import joinedload
import redis

load_dotenv()

from database import db, migrate, Resume, Application, ScrapedJD

def make_celery(app):
    celery = Celery(
        app.import_name,
        backend=app.config['result_backend'],
        broker=app.config['broker_url']
    )
    celery.conf.update(app.config)

    class ContextTask(Task):
        def __call__(self, *args, **kwargs):
            with app.app_context():
                return self.run(*args, **kwargs)

    celery.Task = ContextTask
    return celery

app = Flask(__name__)
CORS(app)

app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'your-secret-key-here')
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///resumeai.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['broker_url'] = 'redis://localhost:6379/0'
app.config['result_backend'] = 'redis://localhost:6379/0'

db.init_app(app)
migrate.init_app(app, db)
socketio = SocketIO(app, message_queue='redis://localhost:6379/0', async_mode='eventlet')

# Initialize Redis for caching
redis_client = redis.Redis(host='localhost', port=6379, db=0, decode_responses=True)

celery = make_celery(app)
celery.conf.imports = ('celery_worker',)

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

ALLOWED_EXTENSIONS = {'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def calculate_file_hash(file_path):
    """Calculate SHA-256 hash of a file for cache invalidation"""
    hash_sha256 = hashlib.sha256()
    try:
        with open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()
    except Exception as e:
        print(f"Error calculating file hash: {e}")
        return None

def get_cached_resume_content(file_hash):
    """Get cached resume content from Redis with error handling"""
    if not file_hash:
        return None
    try:
        cached_data = redis_client.get(f"resume_content:{file_hash}")
        if cached_data:
            return json.loads(cached_data)
    except json.JSONDecodeError as e:
        print(f"Error decoding cached JSON for hash {file_hash}: {e}")
        # Delete corrupted cache entry
        try:
            redis_client.delete(f"resume_content:{file_hash}")
        except:
            pass
    except Exception as e:
        print(f"Error retrieving cached content for hash {file_hash}: {e}")
    return None

def set_cached_resume_content(file_hash, content):
    """Cache resume content in Redis with TTL and error handling"""
    if not file_hash or not content:
        return
    try:
        redis_client.setex(
            f"resume_content:{file_hash}",
            7200,  # 2 hour TTL (increased from 1 hour)
            json.dumps(content, ensure_ascii=False)
        )
    except Exception as e:
        print(f"Error caching content for hash {file_hash}: {e}")

def clear_resume_cache(file_hash):
    """Clear cached resume content from Redis"""
    if not file_hash:
        return
    try:
        redis_client.delete(f"resume_content:{file_hash}")
    except Exception as e:
        print(f"Error clearing cache for hash {file_hash}: {e}")

class ResumeProcessor:
    def __init__(self):
        self.gemini_models = {
            'gemini-2.5-flash': 'gemini-2.5-flash',
            'gemini-2.5-pro': 'gemini-2.5-pro'
        }

        # Load multiple API keys from environment
        self.api_keys = self._load_api_keys()

    def _load_api_keys(self):
        """Load multiple API keys from environment variables"""
        api_keys = []

        # Try to load comma-separated API keys first
        combined_keys = os.environ.get('GEMINI_API_KEYS', '')
        if combined_keys:
            api_keys = [key.strip() for key in combined_keys.split(',') if key.strip()]

        # If no combined keys, try individual keys
        if not api_keys:
            for i in range(1, 11):  # Support up to 10 individual keys
                key = os.environ.get(f'GEMINI_API_KEY_{i}')
                if key:
                    api_keys.append(key)
                else:
                    break

        # Fallback to single key for backward compatibility
        if not api_keys:
            single_key = os.environ.get('GEMINI_API_KEY')
            if single_key:
                api_keys = [single_key]

        return api_keys

    def get_next_api_key(self, used_keys=None):
        """Get the next available API key, rotating through available keys"""
        if used_keys is None:
            used_keys = set()

        available_keys = [key for key in self.api_keys if key not in used_keys]

        if not available_keys:
            if used_keys:
                # Reset and try all keys again
                available_keys = self.api_keys
            else:
                raise Exception("No API keys available")

        return available_keys[0]

    def _call_gemini_api_with_fallback(self, model, prompt, request_options=None, used_keys=None):
        """Call Gemini API with automatic fallback to other keys and models"""
        if used_keys is None:
            used_keys = set()

        if request_options is None:
            request_options = {}

        last_error = None

        # Try current model with all available API keys
        while len(used_keys) < len(self.api_keys):
            try:
                api_key = self.get_next_api_key(used_keys)
                used_keys.add(api_key)

                print(f"Trying API key {len(used_keys)}/{len(self.api_keys)} with model {model}")
                genai.configure(api_key=api_key)
                model_instance = genai.GenerativeModel(self.gemini_models[model])
                response = model_instance.generate_content(prompt, request_options=request_options)

                # Validate response
                if response and response.text:
                    json_match = re.search(r'\{.*\}', response.text, re.DOTALL)
                    if not json_match:
                        if response.text.strip().startswith('{'):
                            return json.loads(response.text, strict=False)
                        raise Exception("AI response did not contain a valid JSON object.")
                    return json.loads(json_match.group(), strict=False)
                else:
                    raise Exception("Empty response from AI model")

            except Exception as e:
                print(f"API key {len(used_keys)} failed: {str(e)}")
                last_error = e
                continue

        # If all API keys failed for current model, try alternative model
        alternative_model = 'gemini-2.5-pro' if model == 'gemini-2.5-flash' else 'gemini-2.5-flash'

        if alternative_model != model:
            print(f"All API keys failed for {model}, trying {alternative_model}")
            try:
                return self._call_gemini_api_with_fallback(alternative_model, prompt, request_options, set())
            except Exception as e:
                last_error = e

        # If everything failed, raise the last error
        if last_error:
            raise last_error
        else:
            raise Exception("All API keys and models failed")

    def extract_text_from_docx(self, file_path):
        try:
            doc = Document(file_path)
            paragraphs = []
            full_text = []
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if not text:
                    continue
                full_text.append(text)
                para_info = {'id': i, 'text': text}
                paragraphs.append(para_info)
            return {'paragraphs': paragraphs, 'full_text': '\n'.join(full_text)}
        except Exception as e:
            raise Exception(f"Error processing DOCX file: {str(e)}")

    def _call_gemini_api(self, model, prompt, request_options=None):
        if request_options is None:
            request_options = {}
        try:
            response = model.generate_content(prompt, request_options=request_options)

            if not response or not response.text:
                raise Exception("Empty response from AI model")

            response_text = response.text.strip()

            # Log full response for debugging (truncate if too long)
            if len(response_text) > 1000:
                print(f"AI Response (first 1000 chars of {len(response_text)} total): {response_text[:1000]}")
                print(f"AI Response (last 1000 chars): {response_text[-1000:]}")
            else:
                print(f"AI Response (full {len(response_text)} chars): {response_text}")

            # Try to extract JSON from the response
            json_match = re.search(r'\{.*\}', response_text, re.DOTALL)

            if json_match:
                json_str = json_match.group()
                print(f"Extracted JSON (first 1000 chars of {len(json_str)} total): {json_str[:1000]}")

                # Clean the JSON string
                json_str = self._clean_json_string(json_str)

                try:
                    parsed_json = json.loads(json_str, strict=False)
                    print(f"Successfully parsed JSON with keys: {list(parsed_json.keys()) if isinstance(parsed_json, dict) else 'Not a dict'}")
                    return parsed_json
                except json.JSONDecodeError as e:
                    print(f"JSON decode error at line {e.lineno}, column {e.colno}: {e.msg}")
                    print(f"Problematic JSON section: {json_str[max(0, e.pos-100):e.pos+100] if e.pos else 'N/A'}")
                    print(f"Attempting to fix JSON...")

                    # Try to fix common JSON issues
                    json_str = self._fix_common_json_issues(json_str)
                    try:
                        parsed_json = json.loads(json_str, strict=False)
                        print(f"Successfully fixed and parsed JSON with keys: {list(parsed_json.keys()) if isinstance(parsed_json, dict) else 'Not a dict'}")
                        return parsed_json
                    except json.JSONDecodeError as e2:
                        print(f"Failed to fix JSON at line {e2.lineno}, column {e2.colno}: {e2.msg}")
                        print(f"Problematic section after fix: {json_str[max(0, e2.pos-100):e2.pos+100] if e2.pos else 'N/A'}")
                        # If all else fails, try to extract just the essential parts
                        return self._extract_json_fallback(response_text)

            # If no JSON found but response starts with {, try the whole response
            if response_text.startswith('{'):
                json_str = self._clean_json_string(response_text)
                try:
                    parsed_json = json.loads(json_str, strict=False)
                    print(f"Successfully parsed full response JSON with keys: {list(parsed_json.keys()) if isinstance(parsed_json, dict) else 'Not a dict'}")
                    return parsed_json
                except json.JSONDecodeError as e:
                    print(f"JSON decode error on full response at line {e.lineno}, column {e.colno}: {e.msg}")
                    print(f"Problematic section: {json_str[max(0, e.pos-100):e.pos+100] if e.pos else 'N/A'}")
                    json_str = self._fix_common_json_issues(json_str)
                    try:
                        parsed_json = json.loads(json_str, strict=False)
                        print(f"Successfully fixed and parsed full response JSON with keys: {list(parsed_json.keys()) if isinstance(parsed_json, dict) else 'Not a dict'}")
                        return parsed_json
                    except json.JSONDecodeError as e2:
                        print(f"Failed to fix JSON on full response at line {e2.lineno}, column {e2.colno}: {e2.msg}")
                        print(f"Problematic section after fix: {json_str[max(0, e2.pos-100):e2.pos+100] if e2.pos else 'N/A'}")
                        return self._extract_interview_prep_fallback(response_text)

            raise Exception(f"AI response did not contain a valid JSON object. Response length: {len(response_text)} chars. Response preview: {response_text[:500]}...")

        except Exception as e:
            print(f"Error during Gemini API call or JSON parsing: {e}")
            print(f"Full response text was: {response.text if 'response' in locals() else 'N/A'}")
            raise

    def _clean_json_string(self, json_str):
        """Clean common JSON formatting issues"""
        # Remove any markdown code blocks
        json_str = re.sub(r'```json\s*', '', json_str)
        json_str = re.sub(r'```\s*', '', json_str)

        # Fix trailing commas before closing braces/brackets
        json_str = re.sub(r',(\s*[}\]])', r'\1', json_str)

        # More robust quote escaping - handle quotes within string values
        # This is more complex because we need to be careful not to break valid JSON
        json_str = self._escape_quotes_in_json(json_str)

        # Additional cleaning for common issues in interview prep responses
        # Remove any text before the first opening brace and after the last closing brace
        first_brace = json_str.find('{')
        last_brace = json_str.rfind('}')
        if first_brace != -1 and last_brace != -1 and last_brace > first_brace:
            json_str = json_str[first_brace:last_brace+1]

        return json_str.strip()

    def _escape_quotes_in_json(self, json_str):
        """More robust quote escaping for JSON strings"""
        try:
            # Parse the JSON to identify string values that need escaping
            # We'll use a more targeted approach to avoid breaking valid JSON

            # First, let's try to identify and fix common issues
            lines = json_str.split('\n')
            cleaned_lines = []

            for line in lines:
                # Skip lines that are just structural (keys, braces, etc.)
                stripped = line.strip()
                if not stripped or stripped in ['{', '}', '[', ']', ','] or stripped.startswith('"') and stripped.endswith('",') or stripped.endswith('":'):
                    cleaned_lines.append(line)
                    continue

                # For lines that contain actual content, escape problematic characters
                if '"' in line:
                    # Find the content between the first and last quote on this line
                    first_quote = line.find('"')
                    last_quote = line.rfind('"')

                    if first_quote != -1 and last_quote != -1 and first_quote != last_quote:
                        # Extract the part before first quote, the content, and after last quote
                        before = line[:first_quote + 1]  # Include the opening quote
                        content = line[first_quote + 1:last_quote]
                        after = line[last_quote:]  # Include the closing quote and rest

                        # Escape quotes and backslashes in the content
                        escaped_content = content.replace('\\', '\\\\').replace('"', '\\"')
                        cleaned_lines.append(before + escaped_content + after)
                    else:
                        cleaned_lines.append(line)
                else:
                    cleaned_lines.append(line)

            return '\n'.join(cleaned_lines)

        except Exception as e:
            print(f"Error in quote escaping: {e}")
            # Fallback to simple replacement
            return json_str.replace('\\"', '"').replace('\\\\', '\\')

    def _fix_common_json_issues(self, json_str):
        """Fix common JSON formatting issues"""
        # Fix trailing commas
        json_str = re.sub(r',(\s*[}\]])', r'\1', json_str)

        # Fix missing commas between objects
        json_str = re.sub(r'}(\s*){', r'},', json_str)

        # Fix missing commas between arrays
        json_str = re.sub(r'](\s*)(\[)', r'],\1\2', json_str)

        # Fix boolean values
        json_str = re.sub(r'\bTrue\b', 'true', json_str)
        json_str = re.sub(r'\bFalse\b', 'false', json_str)
        json_str = re.sub(r'\bNone\b', 'null', json_str)

        # Fix unescaped quotes in string values - this is the main issue
        json_str = self._fix_unescaped_quotes(json_str)

        return json_str

    def _fix_unescaped_quotes(self, json_str):
        """Fix unescaped quotes within JSON string values"""
        try:
            # More aggressive approach: find all string values and escape them properly
            # This regex finds string values (content between quotes) while avoiding key names
            import re

            # Pattern to match string values (not keys) - looks for content after colon and comma
            # This handles cases like: "key": "value with unescaped 'quotes' in it"
            string_pattern = r'("[^"]*")\s*:\s*"((?:[^"\\]|\\.)*")'

            def escape_string_content(match):
                key = match.group(1)  # The key part (e.g., "answer":)
                content = match.group(2)  # The content part

                # Remove the trailing quote to work with just the content
                if content.endswith('"'):
                    content = content[:-1]

                # Escape problematic characters in the content
                # First, replace any existing backslashes to avoid double escaping
                escaped_content = content.replace('\\', '\\\\')
                # Then escape quotes
                escaped_content = escaped_content.replace('"', '\\"')
                # Handle newlines, tabs, and other special characters
                escaped_content = escaped_content.replace('\n', '\\n').replace('\r', '\\r').replace('\t', '\\t')

                return f'{key}: "{escaped_content}"'

            # Apply the fix to key-value pairs
            fixed_json = re.sub(string_pattern, escape_string_content, json_str)

            # Also handle string values in arrays (like in talking_points)
            # Pattern to match array elements: "item1", "item with 'quotes'", "item3"
            array_string_pattern = r',\s*"((?:[^"\\]|\\.)*")(?=\s*[,\]])'
            single_array_item_pattern = r'\[\s*"((?:[^"\\]|\\.)*)"(?=\s*\])'  # For single item in array

            def escape_array_item(match):
                content = match.group(1)

                # Remove the trailing quote to work with just the content
                if content.endswith('"'):
                    content = content[:-1]

                # Escape problematic characters in the content
                escaped_content = content.replace('\\', '\\\\').replace('"', '\\"').replace('\n', '\\n').replace('\r', '\\r').replace('\t', '\\t')

                return f', "{escaped_content}"'

            def escape_single_array_item(match):
                content = match.group(1)

                # Remove the trailing quote to work with just the content
                if content.endswith('"'):
                    content = content[:-1]

                # Escape problematic characters in the content
                escaped_content = content.replace('\\', '\\\\').replace('"', '\\"').replace('\n', '\\n').replace('\r', '\\r').replace('\t', '\\t')

                return f'["{escaped_content}"'

            # Apply fixes to array elements
            fixed_json = re.sub(array_string_pattern, escape_array_item, fixed_json)
            fixed_json = re.sub(single_array_item_pattern, escape_single_array_item, fixed_json)

            # Handle remaining standalone strings that might be unescaped
            # This is more complex - we'll use a state-based approach to identify string values vs. keys
            result = []
            i = 0
            in_string = False
            string_start_char = None
            escape_next = False

            while i < len(fixed_json):
                char = fixed_json[i]

                if escape_next:
                    result.append(char)
                    escape_next = False
                elif char == '\\':
                    result.append(char)
                    escape_next = True
                elif char in ('"', "'") and not escape_next:
                    if not in_string:
                        # Starting a new string
                        in_string = True
                        string_start_char = char
                        result.append(char)
                    elif char == string_start_char:
                        # Ending the current string
                        in_string = False
                        string_start_char = None
                        result.append(char)
                    else:
                        # This is a different quote character inside a string, escape it
                        result.append('\\')
                        result.append(char)
                else:
                    result.append(char)

                i += 1

            return ''.join(result)

        except Exception as e:
            print(f"Error in quote fixing: {e}")
            # Fallback: try to manually escape all quotes that aren't already escaped
            try:
                # Simple fallback: escape all unescaped quotes
                result = []
                i = 0
                while i < len(json_str):
                    if json_str[i] == '"' and (i == 0 or json_str[i-1] != '\\'):
                        # This is an unescaped quote, escape it
                        result.append('\\"')
                    else:
                        result.append(json_str[i])
                    i += 1
                return ''.join(result)
            except:
                return json_str

    def _extract_json_fallback(self, response_text):
        """Fallback method to extract JSON when parsing fails"""
        print(f"Attempting fallback JSON extraction from response text...")

        # Try to extract cover letter
        cover_letter_match = re.search(r'"cover_letter"\s*:\s*"([^"]*(?:\\.[^"]*)*)"', response_text, re.DOTALL)
        if cover_letter_match:
            cover_letter = cover_letter_match.group(1)
            # Unescape the content
            cover_letter = cover_letter.replace('\\"', '"').replace('\\\\', '\\').replace('\\n', '\n')
        else:
            cover_letter = 'Error: Could not parse AI response'

        # Try to extract match score
        match_score_match = re.search(r'"match_score"\s*:\s*(\d+)', response_text)
        match_score = int(match_score_match.group(1)) if match_score_match else 0

        # Try to extract match score analysis
        analysis_match = re.search(r'"match_score_analysis"\s*:\s*({[^}]*})', response_text, re.DOTALL)
        if analysis_match:
            try:
                analysis = json.loads(analysis_match.group(1))
            except:
                analysis = {
                    'strengths': 'Error parsing analysis',
                    'gaps': 'Error parsing analysis',
                    'justification': 'Error parsing analysis'
                }
        else:
            analysis = {
                'strengths': 'Error parsing analysis',
                'gaps': 'Error parsing analysis',
                'justification': 'Error parsing analysis'
            }

        # Try to extract paragraphs if present
        paragraphs = {}
        if 'customized_paragraphs' in response_text:
            para_matches = re.findall(r'"([^"]+)"\s*:\s*"([^"]*(?:\\.[^"]*)*)"', response_text)
            for key, value in para_matches:
                if key.startswith('paragraph') or len(key) < 50:  # Likely a paragraph key
                    # Unescape the value
                    unescaped_value = value.replace('\\"', '"').replace('\\\\', '\\').replace('\\n', '\n')
                    paragraphs[key] = unescaped_value

        result = {
            'customized_paragraphs': paragraphs,
            'cover_letter': cover_letter,
            'match_score': match_score,
            'match_score_analysis': analysis
        }

        print(f"Fallback extraction successful: {bool(cover_letter != 'Error: Could not parse AI response')}")
        return result

    def _extract_interview_prep_fallback(self, response_text):
        """Specialized fallback method for interview prep JSON parsing"""
        print(f"Attempting specialized interview prep JSON extraction...")

        # Try to extract general_questions array
        general_questions = []
        role_based_questions = []

        # Look for general_questions array
        general_match = re.search(r'"general_questions"\s*:\s*(\[[^\]]*\])', response_text, re.DOTALL)
        if general_match:
            try:
                general_questions = json.loads(general_match.group(1))
                print(f"Successfully extracted {len(general_questions)} general questions")
            except json.JSONDecodeError as e:
                print(f"Failed to parse general_questions array: {e}")

        # Look for role_based_questions array
        role_match = re.search(r'"role_based_questions"\s*:\s*(\[[^\]]*\])', response_text, re.DOTALL)
        if role_match:
            try:
                role_based_questions = json.loads(role_match.group(1))
                print(f"Successfully extracted {len(role_based_questions)} role-based questions")
            except json.JSONDecodeError as e:
                print(f"Failed to parse role_based_questions array: {e}")

        # If we got at least one type of questions, return them
        if general_questions or role_based_questions:
            result = {
                'general_questions': general_questions,
                'role_based_questions': role_based_questions
            }
            print(f"Interview prep fallback extraction successful: {len(general_questions)} general, {len(role_based_questions)} role-based questions")
            return result

        # If that didn't work, try a more aggressive approach
        print("Trying aggressive JSON extraction for interview prep...")

        # Look for any arrays that might contain question objects
        question_arrays = re.findall(r'\[\s*\{[^}]*"question"[^}]*\}[^\]]*\]', response_text, re.DOTALL)

        for i, array_str in enumerate(question_arrays):
            try:
                questions = json.loads(array_str)
                if isinstance(questions, list) and len(questions) > 0:
                    if i == 0:
                        general_questions = questions
                        print(f"Extracted {len(general_questions)} general questions from array {i}")
                    else:
                        role_based_questions = questions
                        print(f"Extracted {len(role_based_questions)} role-based questions from array {i}")
            except json.JSONDecodeError as e:
                print(f"Failed to parse question array {i}: {e}")

        if general_questions or role_based_questions:
            result = {
                'general_questions': general_questions,
                'role_based_questions': role_based_questions
            }
            print(f"Aggressive interview prep extraction successful: {len(general_questions)} general, {len(role_based_questions)} role-based questions")
            return result

        print("All interview prep extraction methods failed")
        return {
            'general_questions': [],
            'role_based_questions': []
        }

    # MODIFIED: Logic to handle custom prompts
    def _get_prompt(self, prompt_key, custom_prompts_dict, placeholders):
        default_prompts = {
            "paragraphs": """ROLE:
You are an elite Career Strategist and Certified Professional Resume Writer (CPRW) with deep expertise in Applicant Tracking System (ATS) optimization and modern recruitment psychology. Your specialization is reverse-engineering job descriptions to create compelling career narratives that bypass algorithmic filters and resonate with hiring managers at top-tier companies like {COMPANY}.

OBJECTIVE:
Your mission is to strategically re-engineer the provided resume paragraphs. Transform them from passive descriptions of duties into high-impact, quantified statements of achievement. The rewritten paragraphs must be meticulously tailored to the target job description, demonstrating an undeniable fit for the role.

CONTEXTUAL INPUTS:

COMPANY: {COMPANY}

TARGET JOB DESCRIPTION: {JOB_DESCRIPTION}

PARAGRAPHS FOR TRANSFORMATION ({PARAGRAPH_COUNT} total): {SELECTED_PARAGRAPHS_JSON}

MAXIMUM TOTAL CHARACTER COUNT: {TOTAL_CHAR_LIMIT}

EXECUTION DIRECTIVES:

ATS & Keyword Optimization (Primary Directive):

Analyze & Map: First, meticulously parse the {JOB_DESCRIPTION} to identify primary and secondary keywords. This includes hard skills (e.g., software, technical methodologies), soft skills (e.g., 'strategic planning', 'cross-functional collaboration'), and key qualifications.

Semantic Integration: Do not merely "stuff" keywords. Integrate them naturally and semantically. If the JD mentions "managing budgets," use related powerful phrases like "financial oversight," "P&L management," or "resource allocation" if supported by the original text.

Mirror Language: Reflect the specific terminology and professional tone used by {COMPANY} in the job description to create a sense of immediate cultural and professional alignment.

Quantification & Impact Framing (Secondary Directive):

Employ the STAR/PAR Method: Restructure every possible statement to follow the Problem-Action-Result (or Situation-Task-Action-Result) framework. Focus on the outcome of the actions.

Introduce Metrics: Where the original text implies an achievement, quantify it. Use metrics such as percentages (e.g., increased efficiency by 15%), monetary values (e.g., managed a £500K budget), scale (e.g., led a team of 10), or time saved (e.g., reduced processing time by 2 days). The goal is to translate responsibilities into measurable results.

Lead with Impact: Begin sentences with a powerful, diverse action verb that immediately signals achievement (e.g., "Orchestrated," "Engineered," "Spearheaded," "Maximized," "Revitalized"). Avoid passive language ("Responsible for...") and low-impact verbs ("Led," "Managed") where a stronger alternative exists.

Structural & Stylistic Integrity:

Conciseness: Eliminate filler words and redundant phrases. Each word must serve a purpose.

High-Fidelity Transformation: You must adhere strictly to the achievements and experiences present in the original {SELECTED_PARAGRAPHS_JSON}. Enhance and reframe, but never fabricate new data, skills, or outcomes.

Adhere to Constraints: The combined character count of all transformed paragraphs must not exceed the {TOTAL_CHAR_LIMIT}. The output must be a direct one-to-one transformation of the provided paragraph IDs.

CRITICAL FINAL CHECK:
Before finalizing, review the rewritten paragraphs against the {JOB_DESCRIPTION} one last time. Ask: "Does this text make the candidate look like the perfect solution to the problems and needs outlined in this job description?" The answer must be an unequivocal "yes."
""",
            "single_paragraph": """ROLE:
You are an elite Career Strategist and Certified Professional Resume Writer (CPRW) with deep expertise in Applicant Tracking System (ATS) optimization and modern recruitment psychology. Your specialization is reverse-engineering job descriptions to create compelling career narratives that bypass algorithmic filters and resonate with hiring managers at top-tier companies like {COMPANY}.

OBJECTIVE:
Your mission is to strategically re-engineer a single resume paragraph. Transform it from a passive description of duties into a high-impact, quantified statement of achievement. The rewritten paragraph must be meticulously tailored to the target job description, demonstrating an undeniable fit for the role.

CONTEXTUAL INPUTS:

COMPANY: {COMPANY}

TARGET JOB DESCRIPTION: {JOB_DESCRIPTION}

ORIGINAL PARAGRAPH FOR TRANSFORMATION: "{ORIGINAL_PARAGRAPH}"

MAXIMUM WORD COUNT: {WORD_LIMIT}

EXECUTION DIRECTIVES:

ATS & Keyword Optimization (Primary Directive):

Analyze & Map: First, meticulously parse the {JOB_DESCRIPTION} to identify primary and secondary keywords. This includes hard skills (e.g., software, technical methodologies), soft skills (e.g., 'strategic planning', 'cross-functional collaboration'), and key qualifications.

Semantic Integration: Do not merely "stuff" keywords. Integrate them naturally and semantically. If the JD mentions "managing budgets," use related powerful phrases like "financial oversight," "P&L management," or "resource allocation" if supported by the original text.

Mirror Language: Reflect the specific terminology and professional tone used by {COMPANY} in the job description to create a sense of immediate cultural and professional alignment.

Quantification & Impact Framing (Secondary Directive):

Employ the STAR/PAR Method: Restructure the statement to follow the Problem-Action-Result (or Situation-Task-Action-Result) framework. Focus on the outcome of the actions.

Introduce Metrics: Where the original text implies an achievement, quantify it. Use metrics such as percentages (e.g., increased efficiency by 15%), monetary values (e.g., managed a £500K budget), scale (e.g., led a team of 10), or time saved (e.g., reduced processing time by 2 days). The goal is to translate responsibilities into measurable results.

Lead with Impact: Begin the paragraph with a powerful, diverse action verb that immediately signals achievement (e.g., "Orchestrated," "Engineered," "Spearheaded," "Maximized," "Revitalized"). Avoid passive language ("Responsible for...") and low-impact verbs ("Led," "Managed") where a stronger alternative exists.

Structural & Stylistic Integrity:

Conciseness: Eliminate filler words and redundant phrases. Each word must serve a purpose.

High-Fidelity Transformation: You must adhere strictly to the achievements and experiences present in the original paragraph. Enhance and reframe, but never fabricate new data, skills, or outcomes.

Adhere to Constraints: The transformed paragraph must not exceed {WORD_LIMIT} words. Maintain a similar length to the original while dramatically improving impact and relevance.

CRITICAL FINAL CHECK:
Before finalizing, review the rewritten paragraph against the {JOB_DESCRIPTION} one last time. Ask: "Does this text make the candidate look like the perfect solution to the problems and needs outlined in this job description?" The answer must be an unequivocal "yes."
""",
            "cover_letter": """ROLE: You are an Expert Career Strategist and Recruitment Analyst. Your expertise lies in dissecting job descriptions and resumes to create compelling application materials and provide a rigorous, objective analysis of a candidate's viability. You do not sugarcoat; your feedback is direct, evidence-based, and actionable.

MISSION: Your mission is to perform a two-part task based on the provided context. First, you will write a world-class cover letter that positions the candidate as the ideal solution to the company's needs. Second, you will conduct a brutally honest, data-driven analysis to score the candidate's match for the role, identifying both strengths and critical gaps.

CONTEXT:

COMPANY: {COMPANY}

TARGET JOB DESCRIPTION: {JOB_DESCRIPTION}

FULL RESUME TEXT: {FULL_RESUME_TEXT}

TASK: GENERATE COVER LETTER & STRATEGIC MATCH ANALYSIS

PART 1: THE COVER LETTER (250-300 words)

Your writing must be concise, confident, and meticulously tailored.

Opening Hook: Do not start with a generic "I am writing to apply...". Instead, create a powerful opening statement that immediately connects the candidate's most significant achievement or core competency to a specific company goal, value, or a challenge implied in the job description.

Body Paragraphs (The "Proof"):

Synthesize the top 2-3 requirements from the {JOB_DESCRIPTION}.

For each requirement, extract a specific, quantifiable achievement from the {FULL_RESUME_TEXT} that directly proves the candidate's capability.

Weave these proofs into a narrative. Use the Problem-Action-Result (PAR) framework. For example: "At my previous role, I addressed the challenge of [Problem] by implementing [Action], which resulted in a [Quantifiable Result]."

Subtly integrate knowledge of {COMPANY}'s products, recent news, or mission to demonstrate genuine interest beyond the job posting.

Closing & Call to Action: Conclude with a confident statement summarizing the candidate's value proposition. End with a proactive call to action, expressing eagerness to discuss how their specific skills can contribute to the company's upcoming projects or goals.

PART 2: THE JOB MATCH ANALYSIS

Your analysis must be objective and unflinching. Avoid platitudes.

Job Match Score (1-100): Provide a single integer score based on the following rubric.

90-100 (Exceptional): Candidate exceeds most core requirements, meets all preferred qualifications, and possesses unique value-adds. The resume provides quantifiable proof of high performance in directly comparable tasks.

80-89 (Strong): Candidate meets all core requirements and most preferred qualifications. There is a clear and direct mapping between resume experience and job duties.

70-79 (Good): Candidate meets the majority of core requirements but may be missing some preferred qualifications or lack direct experience in a specific domain. The candidacy is solid but not flawless.

Below 70 (Moderate to Weak): Candidate is missing one or more core requirements. The experience is adjacent or requires significant upskilling. This represents a substantial reach for the candidate.

Match Score Analysis: Provide a detailed rationale for your score, structured in the following three sections:

Strengths of Candidacy: Itemize the strongest points of alignment. Quote specific phrases from the {JOB_DESCRIPTION} and directly map them to accomplishments or skills listed in the {FULL_RESUME_TEXT}.

Potential Gaps / Weaknesses: Identify and explicitly state any significant misalignments. Where does the resume fall short? Note missing technologies, insufficient years of experience in a key area, lack of industry-specific context, or any other core requirement that is not fully substantiated by the resume.

Score Justification: Conclude with a summary paragraph that synthesizes the strengths and weaknesses to explain precisely why the specific score was assigned. For example, "The score of 82 reflects the candidate's exceptional alignment with core technical skills A and B, but is tempered by the lack of direct experience with industry-specific software C, which is listed as a preferred qualification."
""",
            "interview_prep": """You are to act as an elite Tier-1 career coach and interview strategist. Your expertise is in meticulously deconstructing a candidate's history against a target role's requirements to forge a powerful, compelling interview narrative. You do not generate generic questions; you create a bespoke interrogation plan designed to highlight the candidate's unique strengths and proactively address potential weaknesses.

PRIMARY OBJECTIVE:

Analyze the provided {FULL_RESUME_TEXT} in the context of the {JOB_DESCRIPTION} for the {JOB_TITLE} role at {COMPANY}. Your goal is to produce a set of highly targeted interview questions and exemplary answers that will strategically position the candidate for success. The output must be a single, valid JSON object with the exact structure specified below - no additional categories or fields are allowed.

ANALYTICAL FRAMEWORK (Your Internal Process):

Resume-to-JD Synergy and Gap Analysis: First, perform a deep comparison. Identify the top 3-5 areas where the candidate's resume shows exceptional alignment with the job description's core requirements. Conversely, identify any potential "red flags" or gaps—such as a non-traditional career path, a noticeable employment gap, a potential lack of experience in a key area mentioned in the JD, or frequent job changes.

Strategic Narrative Formulation: Based on your analysis, determine the central narrative the candidate should convey. This narrative should be woven through all the answers. For example, is it a story of "the technical expert pivoting to leadership," "the generalist now specializing," or "the problem-solver who thrives in chaotic environments"?

TASK: GENERATE INTERVIEW QUESTIONS & ANSWERS

Based on your analysis, generate two distinct categories of questions. For each question, provide both 'talking_points' (the strategic pillars of the response) and a complete sample 'answer' (a polished, first-person narrative).

Category 1: General & Career Narrative Questions (2 Questions)

Mandate: These questions must stem directly from your analysis of the candidate's career trajectory as presented in the resume. They should probe their motivations, rationale for key transitions, and self-awareness. Target the potential "red flags" you identified, framing them as opportunities for the candidate to demonstrate growth, resilience, or strategic thinking. Do not ask generic questions like "Tell me about yourself." Instead, ask pointed questions like, "I noticed you transitioned from [Industry/Role A] to [Industry/Role B]. What catalyzed that specific change, and how did it prepare you for the challenges outlined in our job description?"

Answer Construction: The answers here should solidify the candidate's career narrative. They must explain the "why" behind their decisions, connecting past experiences to their future ambitions for this specific role.

Category 2: Role-Based Questions (2 Questions)

Mandate: These questions must be surgical strikes that connect a specific, critical requirement from the {JOB_DESCRIPTION} with a concrete project or achievement from the {FULL_RESUME_TEXT}. Frame the questions behaviorally to compel storytelling. For example, instead of "Do you have experience with X?", ask, "The job requires extensive experience with [Tool/Skill X from JD]. Describe your most challenging project from your time at [Company from Resume] where you leveraged this skill to overcome a significant obstacle."

Answer Construction: The answers MUST implicitly or explicitly follow the STAR method (Situation, Task, Action, Result).

Situation: Briefly set the context of the project or challenge.

Task: Describe the specific goal or objective you were responsible for.

Action: Detail the specific, individual steps you took to address the task. This is the most important part.

Result: Quantify the outcome. Use metrics, data, and tangible business impact (e.g., "reduced latency by 30%", "increased user engagement by 15%", "saved the project $50k in operational costs"). The result must tie back to the value sought in the job description.

STRICT OUTPUT REQUIREMENTS:

You must return ONLY a valid JSON object with the following EXACT structure. Do not include any text before or after the JSON object. Do not create additional categories beyond these two (no "behavioral_questions" or other categories).

IMPORTANT: Ensure that ALL quotes within your JSON string values are properly escaped with backslashes (e.g., "He said \"Hello\" to me"). Also ensure that all special characters like newlines are properly escaped as \\n. This is CRITICAL for the JSON to be parseable.

```json
{
  "general_questions": [
    {
      "question": "Specific question targeting career narrative...",
      "talking_points": ["Key point 1", "Key point 2", "Key point 3"],
      "answer": "Complete first-person answer using STAR method with properly escaped quotes like \\\"example\\\"..."
    }
  ],
  "role_based_questions": [
    {
      "question": "Specific role-based question...",
      "talking_points": ["Key point 1", "Key point 2", "Key point 3"],
      "answer": "Complete first-person answer using STAR method with properly escaped quotes like \\\"example\\\"..."
    }
  ]
}
```

QUALITY DIRECTIVES:

No Generic Content: Every question and answer must be rigorously tailored to the provided resume and job description.

Strategic Talking Points: The talking_points should not be a mere summary of the answer. They should be concise, strategic bullet points outlining the core message and the key skills being demonstrated (e.g., "Demonstrate proactive problem-solving," "Highlight quantitative impact," "Connect past project to this company's specific needs").

Authentic Voice: The sample answer should be written in a confident, professional, and natural first-person voice. It should be comprehensive but not verbose."""
        }

        # Get the base prompt (custom or default)
        prompt_template = (custom_prompts_dict or {}).get(prompt_key) or default_prompts.get(prompt_key)

        # Replace placeholders
        for key, value in placeholders.items():
            placeholder_tag = f"{{{key}}}"
            prompt_template = prompt_template.replace(placeholder_tag, str(value))

        # ALWAYS append JSON structure requirement, even for custom prompts
        if prompt_key == 'paragraphs':
            json_requirement = "\n\nCRITICAL OUTPUT: Your entire response MUST be a single, valid JSON object with this exact structure:\n" + placeholders.get('JSON_STRUCTURE', '{ "customized_paragraphs": { "paragraph_id_1": "new_text_1", ... } }')
        elif prompt_key == 'single_paragraph':
            json_requirement = "\n\nCRITICAL OUTPUT: Your entire response MUST be a single, valid JSON object with this exact structure:\n" + placeholders.get('JSON_STRUCTURE', '{ "enhanced_text": "The new, enhanced paragraph text here..." }')
        elif prompt_key == 'cover_letter':
            json_requirement = "\n\nCRITICAL OUTPUT: Your entire response MUST be a single, valid JSON object with this exact structure:\n" + placeholders.get('JSON_STRUCTURE', '{\n "cover_letter": "The full cover letter text here...",\n  "match_score": 85,\n  "match_score_analysis": {\n    "strengths": "Strengths of candidacy...",\n    "gaps": "Potential gaps and weaknesses...",\n    "justification": "Score justification..."\n }\n}')
        elif prompt_key == 'interview_prep':
            json_requirement = "\n\nCRITICAL OUTPUT: Your entire response MUST be a single, valid JSON object with this exact structure. Do not add any text or markdown before or after the JSON object.\n" + placeholders.get('JSON_STRUCTURE', '{\n "general_questions": [\n    { "question": "...", "talking_points": ["..."], "answer": "..." }\n ],\n  "role_based_questions": [\n    { "question": "...", "talking_points": ["..."], "answer": "..." }\n  ]\n}')

        prompt_template += json_requirement

   

        return prompt_template

    def _generate_paragraphs(self, model, resume_data, selected_paragraph_ids, job_description, company_name, regenerate_type, custom_prompts):
        if isinstance(regenerate_type, dict) and 'single_paragraph' in regenerate_type:
            para_text = regenerate_type['single_paragraph']
            original_words = len(para_text.split())
            placeholders = {
                'COMPANY': company_name,
                'JOB_DESCRIPTION': job_description,
                'ORIGINAL_PARAGRAPH': para_text,
                'WORD_LIMIT': original_words + 15,
                'JSON_STRUCTURE': '{ "enhanced_text": "The new, enhanced paragraph text here..." }'
            }
            prompt = self._get_prompt('single_paragraph', custom_prompts, placeholders)
        else:
            selected_paragraphs_dict = {p['id']: p['text'] for p in resume_data['paragraphs'] if p['id'] in selected_paragraph_ids}
            total_original_words = sum(len(text.split()) for text in selected_paragraphs_dict.values())
            placeholders = {
                'COMPANY': company_name,
                'JOB_DESCRIPTION': job_description,
                'PARAGRAPH_COUNT': len(selected_paragraphs_dict),
                'SELECTED_PARAGRAPHS_JSON': json.dumps(selected_paragraphs_dict, indent=2),
                'TOTAL_WORD_LIMIT': total_original_words + 20,
                'JSON_STRUCTURE': '{ "customized_paragraphs": { "paragraph_id_1": "new_text_1", ... } }'
            }
            prompt = self._get_prompt('paragraphs', custom_prompts, placeholders)
        
        return self._call_gemini_api(model, prompt)

    def _reconstruct_resume_with_enhanced_paragraphs(self, resume_data, enhanced_paragraphs):
        """
        Reconstruct the full resume text by replacing enhanced paragraphs while preserving
        all other content, formatting, and structure.
        """
        if not enhanced_paragraphs:
            return resume_data

        # Create a mapping of original text to enhanced text
        text_replacements = {original: enhanced for original, enhanced in enhanced_paragraphs.items()}

        # Split the full text into lines to preserve structure
        lines = resume_data['full_text'].split('\n')
        reconstructed_lines = []

        for line in lines:
            original_line = line.strip()
            # Check if this line contains any of our paragraphs to replace
            for original_text, enhanced_text in text_replacements.items():
                if original_text.strip() in original_line:
                    # Replace the paragraph content while preserving indentation/formatting
                    if original_line == original_text.strip():
                        # Exact match - replace entire line
                        reconstructed_lines.append(line.replace(original_text, enhanced_text))
                    else:
                        # Partial match - try to replace just the paragraph portion
                        # This handles cases where paragraphs might have slight formatting differences
                        reconstructed_lines.append(line.replace(original_text, enhanced_text))
                    break
            else:
                # No replacement needed for this line
                reconstructed_lines.append(line)

        # Reconstruct the resume data with enhanced text
        enhanced_resume_data = resume_data.copy()
        enhanced_resume_data['full_text'] = '\n'.join(reconstructed_lines)

        return enhanced_resume_data

    def _generate_cover_letter(self, model, resume_data, job_description, company_name, custom_prompts):
        placeholders = {
            'COMPANY': company_name,
            'JOB_DESCRIPTION': job_description,
            'FULL_RESUME_TEXT': resume_data['full_text'],  # Now uses enhanced text
            'JSON_STRUCTURE': '{\n  "cover_letter": "The full cover letter text here...",\n  "match_score": 85,\n  "match_score_analysis": {\n    "strengths": "Strengths of candidacy...",\n    "gaps": "Potential gaps and weaknesses...",\n    "justification": "Score justification..."\n  }\n}'
        }
        prompt = self._get_prompt('cover_letter', custom_prompts, placeholders)
        return self._call_gemini_api(model, prompt)

    def generate_ai_customization(self, api_key, model_name, resume_data, selected_paragraph_ids, job_description, company_name, regenerate_type=None, custom_prompts=None):
        try:
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel(self.gemini_models[model_name])
            
            final_output = {'customized_paragraphs': {}, 'cover_letter': '', 'match_score': None, 'enhanced_text': None}

            do_paragraphs = regenerate_type is None or regenerate_type == 'paragraphs' or isinstance(regenerate_type, dict)
            do_cover_letter = regenerate_type is None or regenerate_type == 'cover_letter'

            if do_paragraphs:
                para_result = self._generate_paragraphs(model, resume_data, selected_paragraph_ids, job_description, company_name, regenerate_type, custom_prompts)
                final_output['enhanced_text'] = para_result.get('enhanced_text')
                if 'customized_paragraphs' in para_result:
                    id_to_text_map = {p['id']: p['text'] for p in resume_data['paragraphs']}
                    for pid, enhanced_text in para_result['customized_paragraphs'].items():
                        try:
                            original_text = id_to_text_map[int(pid)]
                            final_output['customized_paragraphs'][original_text] = enhanced_text
                        except (KeyError, ValueError):
                            print(f"!! DEBUG WARNING: AI returned paragraph ID '{pid}' which was not found. Skipping.")

            if do_cover_letter:
                # NEW: Reconstruct resume text with enhanced paragraphs before generating cover letter
                enhanced_resume_data = self._reconstruct_resume_with_enhanced_paragraphs(resume_data, final_output['customized_paragraphs'])
                cl_result = self._generate_cover_letter(model, enhanced_resume_data, job_description, company_name, custom_prompts)
                final_output['cover_letter'] = cl_result.get('cover_letter')
                final_output['match_score'] = cl_result.get('match_score')
                # Handle both old string format and new structured format for backward compatibility
                match_analysis = cl_result.get('match_score_analysis')
                if isinstance(match_analysis, dict):
                    # New structured format
                    final_output['match_score_analysis'] = match_analysis
                else:
                    # Old string format - convert to structured for consistency
                    final_output['match_score_analysis'] = {
                        'strengths': match_analysis or '',
                        'gaps': '',
                        'justification': 'Analysis converted from legacy format'
                    }

            return final_output
        except Exception as e:
            print(f"AI Generation Error: {traceback.format_exc()}")
            raise Exception(f"Error generating AI customization: {str(e)}")
            
    def generate_interview_prep(self, api_key, model_name, resume_full_text, job_description, company_name, job_title, custom_prompts=None):
        try:
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel(self.gemini_models[model_name])
            
            json_structure = """{
  "general_questions": [
    { "question": "...", "talking_points": ["..."], "answer": "..." }
  ],
  "role_based_questions": [
    { "question": "...", "talking_points": ["..."], "answer": "..." }
  ]
}"""
            placeholders = {
                'COMPANY': company_name,
                'JOB_TITLE': job_title,
                'JOB_DESCRIPTION': job_description,
                'FULL_RESUME_TEXT': resume_full_text,
                'JSON_STRUCTURE': json_structure
            }
            prompt = self._get_prompt('interview_prep', custom_prompts, placeholders)
            return self._call_gemini_api(model, prompt, request_options={"timeout": 300})
        except Exception as e:
            print(f"Interview Prep Generation Error: {traceback.format_exc()}")
            raise Exception(f"Error generating interview prep materials: {str(e)}")

    def update_docx_with_customizations(self, original_file_path, customizations):
        # Handle case where customized_paragraphs might be a JSON string from database
        customized_paragraphs_dict = customizations.get('customized_paragraphs', {})

        # If it's a string (from database), parse it back to dict
        if isinstance(customized_paragraphs_dict, str):
            try:
                customized_paragraphs_dict = json.loads(customized_paragraphs_dict)
            except json.JSONDecodeError:
                print(f"Warning: Could not parse customized_paragraphs JSON string: {customized_paragraphs_dict}")
                customized_paragraphs_dict = {}

        # Ensure it's a dictionary
        if not isinstance(customized_paragraphs_dict, dict):
            print(f"Warning: customized_paragraphs is not a dict: {type(customized_paragraphs_dict)}")
            customized_paragraphs_dict = {}

        try:
            doc = Document(original_file_path)
            for para in doc.paragraphs:
                original_text = para.text.strip()
                if original_text in customized_paragraphs_dict:
                    new_text = customized_paragraphs_dict[original_text]
                    para.text = ""
                    para.add_run(new_text)
            temp_path = tempfile.mktemp(suffix='.docx')
            doc.save(temp_path)
            return temp_path
        except Exception as e:
            raise Exception(f"Error updating DOCX: {str(e)}")

    def create_cover_letter_docx(self, cover_letter_text, company_name, user_name):
        """Create a DOCX document from cover letter text - just plain text, nothing else"""
        try:
            doc = Document()

            # Add cover letter content as plain text - no headers, footers, or formatting
            # Split into paragraphs and add each one
            paragraphs = cover_letter_text.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    para = doc.add_paragraph()
                    para.add_run(para_text.strip())

            # Save to temporary file
            temp_path = tempfile.mktemp(suffix='_resume.docx')
            doc.save(temp_path)
            return temp_path

        except Exception as e:
            raise Exception(f"Error creating cover letter DOCX: {str(e)}")

    def convert_docx_to_pdf(self, docx_path, output_filename):
        """Convert DOCX to PDF using a method that works in Celery workers"""
        try:
            pdf_filename = f"{output_filename}.pdf"
            pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], pdf_filename)

            print(f"Attempting PDF conversion: {docx_path} -> {pdf_path}")

            # Try multiple PDF conversion methods
            conversion_success = False

            # Method 1: Try docx2pdf with better error handling (preserves formatting best)
            try:
                pythoncom.CoInitialize()
                from docx2pdf import convert
                convert(docx_path, pdf_path)

                if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                    print(f"PDF conversion successful using docx2pdf: {pdf_path}")
                    conversion_success = True
                else:
                    raise Exception("PDF file was not created or is empty")

            except Exception as e:
                print(f"docx2pdf failed: {e}")
                # Try method 2: Use pypandoc with DOCX as intermediate format
                try:
                    import pypandoc

                    # First convert DOCX to DOCX (this preserves formatting)
                    # Then convert to PDF
                    temp_docx_path = tempfile.mktemp(suffix='.docx')

                    # Use pandoc to convert DOCX to PDF with better formatting preservation
                    output = pypandoc.convert_file(
                        docx_path,
                        'pdf',
                        outputfile=pdf_path,
                        extra_args=[
                            '--pdf-engine=pdflatex',  # Use LaTeX for better formatting
                            '--standalone',           # Create standalone document
                            '--self-contained',       # Embed all resources
                            '--number-sections',      # Number sections
                            '--toc',                  # Table of contents
                            '--toc-depth=2',          # TOC depth
                        ]
                    )

                    if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                        print(f"PDF conversion successful using pypandoc with LaTeX: {pdf_path}")
                        conversion_success = True
                    else:
                        raise Exception("PDF file was not created or is empty")

                except ImportError:
                    print("pypandoc not available")
                except Exception as e2:
                    print(f"pypandoc failed: {e2}")
                    # Try method 3: Simple pypandoc conversion
                    try:
                        import pypandoc
                        output = pypandoc.convert_file(docx_path, 'pdf', outputfile=pdf_path)
                        if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                            print(f"PDF conversion successful using simple pypandoc: {pdf_path}")
                            conversion_success = True
                        else:
                            raise Exception("PDF file was not created or is empty")
                    except Exception as e3:
                        print(f"Simple pypandoc also failed: {e3}")

            finally:
                # Always clean up COM
                try:
                    pythoncom.CoUninitialize()
                except:
                    pass

            if conversion_success:
                return pdf_path, pdf_filename
            else:
                raise Exception("All PDF conversion methods failed")

        except Exception as e:
            print(f"PDF conversion failed: {e}. Falling back to DOCX.")
            print(f"Error type: {type(e).__name__}")
            print(f"Error details: {str(e)}")

            # Try to clean up any partially created files
            try:
                pdf_filename = f"{output_filename}.pdf"
                pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], pdf_filename)
                if os.path.exists(pdf_path):
                    os.remove(pdf_path)
                    print(f"Cleaned up failed PDF file: {pdf_path}")
            except Exception as cleanup_error:
                print(f"Warning: Could not clean up failed PDF file: {cleanup_error}")

            # Fallback to DOCX
            docx_filename = f"{output_filename}.docx"
            final_docx_path = os.path.join(app.config['UPLOAD_FOLDER'], docx_filename)
            shutil.copy2(docx_path, final_docx_path)
            return final_docx_path, docx_filename

    def convert_docx_to_docx(self, docx_path, output_filename):
        """Simple function to just copy DOCX file - for when user explicitly requests DOCX"""
        try:
            docx_filename = f"{output_filename}.docx"
            final_docx_path = os.path.join(app.config['UPLOAD_FOLDER'], docx_filename)
            shutil.copy2(docx_path, final_docx_path)
            print(f"DOCX copy successful: {final_docx_path}")
            return final_docx_path, docx_filename
        except Exception as e:
            raise Exception(f"Error copying DOCX file: {str(e)}")

processor = ResumeProcessor()

def initialize_system():
    """Initialize system components in background to eliminate first-request delays"""
    import threading
    import time

    def background_init():
        print("🚀 Starting background system initialization...")
        start_time = time.time()

        try:
            # 1. Initialize Python COM for DOCX processing (main culprit for slow first request)
            print("📄 Initializing Python COM...")
            pythoncom.CoInitialize()
            # Create a dummy document to force COM initialization
            try:
                temp_doc = Document()
                temp_doc.add_paragraph("Initialization test")
                temp_path = tempfile.mktemp(suffix='_resume.docx')
                temp_doc.save(temp_path)
                os.remove(temp_path)
                print("✅ COM initialization complete")
            except Exception as e:
                print(f"⚠️ COM initialization warning: {e}")

            # 2. Warm up database connections
            print("🗄️ Warming up database connections...")
            try:
                with app.app_context():
                    # Force database connection initialization
                    db.engine.connect()
                    print("✅ Database connection pool initialized")
            except Exception as e:
                print(f"⚠️ Database initialization warning: {e}")

            # 3. Test Redis connection
            print("🔴 Testing Redis connection...")
            try:
                redis_client.ping()
                print("✅ Redis connection verified")
            except Exception as e:
                print(f"⚠️ Redis connection warning: {e}")

            # 4. Pre-initialize any heavy imports
            print("📦 Pre-loading heavy modules...")
            try:
                # Import modules that might be slow on first use
                import docx2pdf
                print("✅ Heavy modules pre-loaded")
            except Exception as e:
                print(f"⚠️ Module pre-loading warning: {e}")

            elapsed = time.time() - start_time
            print(f"🎉 Background initialization complete in {elapsed:.2f} seconds")

        except Exception as e:
            print(f"❌ Background initialization failed: {e}")
        finally:
            # Clean up COM
            try:
                pythoncom.CoUninitialize()
            except:
                pass

    # Start background initialization in a separate thread
    init_thread = threading.Thread(target=background_init, daemon=True)
    init_thread.start()
    print("🔄 Background initialization started (non-blocking)")

# Initialize system immediately when module loads
initialize_system()

@app.before_request
def before_request_func():
    if 'user_session_id' not in session:
        session['user_session_id'] = str(uuid.uuid4())

@socketio.on('connect')
def handle_connect():
    print(f"SocketIO client connected: {request.sid}")
    if 'user_session_id' in session:
        join_room(session['user_session_id'])
        socketio.emit('session_id', {'id': session['user_session_id']})
        print(f"Joined room: {session['user_session_id']}")

        # Notify client that system is ready (or will be soon)
        # Use a small delay to ensure the client is ready to receive the message
        @socketio.event
        def notify_system_ready():
            socketio.emit('system_status', {
                'status': 'ready',
                'message': 'System initialization complete - Fast downloads enabled!'
            }, room=session['user_session_id'])

        socketio.call_later(0.5, notify_system_ready)

@socketio.on('connect')
def handle_connect():
    print(f"SocketIO client connected: {request.sid}")
    if 'user_session_id' in session:
        join_room(session['user_session_id'])
        socketio.emit('session_id', {'id': session['user_session_id']})
        print(f"Joined room: {session['user_session_id']}")

@socketio.on('join')
def handle_join(data):
    session_id = data.get('session_id')
    if session_id:
        join_room(session_id)
        socketio.emit('session_id', {'id': session_id})
        print(f"Client joined room: {session_id}")

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/customize', methods=['POST'])
def customize_resume():
    try:
        data = request.get_json()
        resume = Resume.query.get_or_404(data.get('resume_id'))
        if resume.user_session_id != session.get('user_session_id'):
            abort(403)
        data['session_id'] = session['user_session_id']
        task = celery.send_task('celery_worker.generate_customization_task', args=[data])
        return jsonify({'job_id': task.id})
    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500
        
@app.route('/save_user_info', methods=['POST'])
def save_user_info():
    data = request.get_json()
    session['user_first_name'] = data.get('first_name', '')
    session['user_last_name'] = data.get('last_name', '')
    return jsonify({'message': 'User information saved successfully'})

@app.route('/api/status', methods=['GET'])
def get_status():
    has_set_name = 'user_first_name' in session and session['user_first_name']
    resumes_count = 0
    if has_set_name:
        resumes_count = Resume.query.filter_by(user_session_id=session.get('user_session_id')).count()
    return jsonify({
        'has_set_name': has_set_name, 
        'resumes_count': resumes_count,
        'session_id': session.get('user_session_id')
    })

@app.route('/api/resumes', methods=['POST'])
def upload_resume():
    if 'resume_file' not in request.files: return jsonify({'error': 'No file part'}), 400
    file = request.files['resume_file']
    if file.filename == '': return jsonify({'error': 'No selected file'}), 400
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{uuid.uuid4()}_{filename}")
        file.save(file_path)

        # Calculate file hash for caching
        file_hash = calculate_file_hash(file_path)

        # Check if we have cached content for this file
        cached_structured_text = get_cached_resume_content(file_hash)

        if not cached_structured_text:
            # Parse the document and cache the result
            try:
                cached_structured_text = processor.extract_text_from_docx(file_path)
                if file_hash:
                    set_cached_resume_content(file_hash, cached_structured_text)
            except Exception as e:
                print(f"Error extracting text from {filename} on upload: {e}")
                cached_structured_text = None

        new_resume = Resume(
            resume_name=request.form.get('resume_name'),
            original_file_path=file_path,
            user_session_id=session.get('user_session_id'),
            user_first_name=request.form.get('first_name'),
            user_last_name=request.form.get('last_name'),
            structured_text=cached_structured_text,
            file_hash=file_hash
        )
        db.session.add(new_resume)
        db.session.commit()
        return jsonify(new_resume.to_dict())
    return jsonify({'error': 'Invalid file type'}), 400

@app.route('/api/resumes', methods=['GET'])
def get_resumes():
    resumes = Resume.query.filter_by(user_session_id=session.get('user_session_id')).all()
    return jsonify([r.to_dict() for r in resumes])

@app.route('/api/resumes/<int:resume_id>', methods=['GET'])
def get_resume_details(resume_id):
    resume = Resume.query.get_or_404(resume_id)
    if resume.user_session_id != session.get('user_session_id'): abort(403)
    details = resume.to_dict()

    try:
        # Priority 1: Use cached data from database if available
        if resume.structured_text and 'paragraphs' in resume.structured_text:
            details['paragraphs'] = resume.structured_text['paragraphs']
        else:
            # Priority 2: Check Redis cache first before parsing
            if resume.file_hash:
                cached_content = get_cached_resume_content(resume.file_hash)
                if cached_content:
                    details['paragraphs'] = cached_content.get('paragraphs', [])
                    # Update the database with cached content for future use
                    resume.structured_text = cached_content
                    db.session.commit()
                else:
                    # Priority 3: Parse once and cache the result
                    resume_content = processor.extract_text_from_docx(resume.original_file_path)
                    details['paragraphs'] = resume_content.get('paragraphs', [])
                    if resume.file_hash:
                        set_cached_resume_content(resume.file_hash, resume_content)
                        resume.structured_text = resume_content
                        db.session.commit()
            else:
                # Priority 4: Fallback to direct parsing if no hash (should be rare)
                resume_content = processor.extract_text_from_docx(resume.original_file_path)
                details['paragraphs'] = resume_content.get('paragraphs', [])
    except Exception as e:
        details['paragraphs'] = []
        details['error'] = f"Could not read DOCX file: {str(e)}"
    return jsonify(details)

@app.route('/api/resumes/<int:resume_id>/selections', methods=['PUT'])
def save_resume_selections(resume_id):
    resume = Resume.query.get_or_404(resume_id)
    if resume.user_session_id != session.get('user_session_id'): abort(403)
    data = request.get_json()
    resume.selected_paragraph_ids = data.get('selected_paragraph_ids', [])
    db.session.commit()
    return jsonify({'message': 'Selections updated successfully'})

@app.route('/api/resumes/<int:resume_id>', methods=['DELETE'])
def delete_resume(resume_id):
    resume = Resume.query.get_or_404(resume_id)
    if resume.user_session_id != session.get('user_session_id'): abort(403)
    try:
        if os.path.exists(resume.original_file_path):
            os.remove(resume.original_file_path)
    except Exception as e:
        print(f"Error deleting file {resume.original_file_path}: {e}")
    db.session.delete(resume)
    db.session.commit()
    return jsonify({'message': 'Resume deleted'})

@app.route('/api/applications', methods=['POST'])
def save_application():
    try:
        data = request.get_json()
        print(f"DEBUG: Saving application with data keys: {list(data.keys()) if data else 'None'}")

        if not data:
            return jsonify({'error': 'No data provided'}), 400

        resume = Resume.query.get_or_404(data.get('resume_id'))

        # Convert match_score_analysis to JSON string if it's an object
        match_score_analysis = data.get('match_score_analysis')
        print(f"DEBUG: match_score_analysis type: {type(match_score_analysis)}")
        print(f"DEBUG: match_score_analysis value: {str(match_score_analysis)[:200]}...")

        if match_score_analysis is not None:
            if isinstance(match_score_analysis, dict):
                print(f"DEBUG: Converting dict to JSON string")
                match_score_analysis = json.dumps(match_score_analysis, ensure_ascii=False)
                print(f"DEBUG: Converted to: {str(match_score_analysis)[:200]}...")
            elif isinstance(match_score_analysis, str):
                print(f"DEBUG: Already a string, keeping as-is")
            else:
                print(f"DEBUG: Converting other type to string")
                match_score_analysis = str(match_score_analysis)

        # Convert customized_paragraphs to JSON string if it's an object
        customized_paragraphs = data.get('customized_paragraphs')
        print(f"DEBUG: customized_paragraphs type: {type(customized_paragraphs)}")

        if customized_paragraphs is not None:
            if isinstance(customized_paragraphs, dict):
                print(f"DEBUG: Converting customized_paragraphs dict to JSON string")
                customized_paragraphs = json.dumps(customized_paragraphs, ensure_ascii=False)
                print(f"DEBUG: Converted customized_paragraphs to: {str(customized_paragraphs)[:200]}...")
            elif isinstance(customized_paragraphs, str):
                print(f"DEBUG: customized_paragraphs already a string, keeping as-is")
            else:
                print(f"DEBUG: Converting customized_paragraphs other type to string")
                customized_paragraphs = str(customized_paragraphs)

        print(f"DEBUG: Final types - match_score_analysis: {type(match_score_analysis)}, customized_paragraphs: {type(customized_paragraphs)}")

        new_app = Application(
            company_name=data.get('company_name'),
            job_title=data.get('job_title'),  # Add job title from scraped JD
            job_description=data.get('job_description', ''),
            status=data.get('status', 'not_applied'),
            match_score=data.get('match_score'),
            match_score_analysis=match_score_analysis,  # Should now be a string
            cover_letter=data.get('cover_letter'),
            customized_paragraphs=customized_paragraphs,  # Should now be a string
            job_posting_url=data.get('job_posting_url'),  # Add job posting URL
            user_session_id=session.get('user_session_id'),
            resume_id=resume.id
        )

        print(f"DEBUG: Created Application object: {new_app.company_name}, {new_app.job_title}")

        db.session.add(new_app)
        db.session.commit()

        print(f"DEBUG: Application saved successfully with ID: {new_app.id}")

        return jsonify(new_app.to_dict())

    except Exception as e:
        print(f"ERROR: Failed to save application: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Failed to save application: {str(e)}'}), 500

@app.route('/api/applications', methods=['GET'])
def get_applications():
    # Use optimized query with eager loading to avoid N+1 queries
    apps = Application.get_with_resume(session.get('user_session_id'))

    # Apply filters based on query parameters
    company_filter = request.args.get('company', '').strip()
    status_filter = request.args.get('status', '').strip()
    min_score = request.args.get('min_score', '').strip()
    max_score = request.args.get('max_score', '').strip()
    from_date = request.args.get('from_date', '').strip()
    to_date = request.args.get('to_date', '').strip()

    # Filter applications
    if company_filter:
        apps = [app for app in apps if company_filter.lower() in app.company_name.lower()]

    if status_filter:
        apps = [app for app in apps if app.status == status_filter]

    if min_score:
        try:
            min_score_val = int(min_score)
            apps = [app for app in apps if app.match_score and app.match_score >= min_score_val]
        except ValueError:
            pass

    if max_score:
        try:
            max_score_val = int(max_score)
            apps = [app for app in apps if app.match_score and app.match_score <= max_score_val]
        except ValueError:
            pass

    if from_date:
        try:
            from_date_obj = datetime.fromisoformat(from_date.replace('Z', '+00:00'))
            apps = [app for app in apps if app.created_date >= from_date_obj]
        except ValueError:
            pass

    if to_date:
        try:
            to_date_obj = datetime.fromisoformat(to_date.replace('Z', '+00:00'))
            apps = [app for app in apps if app.created_date <= to_date_obj]
        except ValueError:
            pass

    return jsonify([app.to_dict() for app in apps])

@app.route('/api/applications/<int:app_id>', methods=['GET'])
def get_application_details(app_id):
    # Use optimized query with eager loading
    app = Application.query.options(joinedload(Application.resume)).filter_by(
        id=app_id,
        user_session_id=session.get('user_session_id')
    ).first_or_404()
    return jsonify(app.to_dict())

@app.route('/api/applications/<int:app_id>', methods=['PUT'])
def update_application(app_id):
    app = Application.query.get_or_404(app_id)
    if app.user_session_id != session.get('user_session_id'): abort(403)
    data = request.get_json()
    if 'status' in data:
        app.status = data['status']
    app.updated_date = datetime.utcnow()
    db.session.commit()
    return jsonify(app.to_dict())

@app.route('/api/applications/<int:app_id>', methods=['DELETE'])
def delete_application(app_id):
    app = Application.query.get_or_404(app_id)
    if app.user_session_id != session.get('user_session_id'): abort(403)
    db.session.delete(app)
    db.session.commit()
    return jsonify({'message': 'Application deleted'})

@app.route('/api/applications/<int:app_id>/generate-interview-prep', methods=['POST'])
def generate_interview_prep(app_id):
    try:
        app = Application.query.get_or_404(app_id)
        if app.user_session_id != session.get('user_session_id'):
            abort(403)
        
        data = request.get_json() or {}
        task_data = {
            'app_id': app.id,
            'session_id': session['user_session_id'],
            'ai_model': data.get('ai_model', 'gemini-2.5-pro'),
            'custom_prompts': data.get('custom_prompts') # Pass custom prompts
        }
        
        task = celery.send_task('celery_worker.generate_interview_prep_task', args=[task_data])
        return jsonify({'job_id': task.id})
    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

# MODIFIED: Endpoint now a POST to receive custom prompt in body
@app.route('/api/applications/<int:app_id>/interview-prompt-text', methods=['POST'])
def get_interview_prompt_text(app_id):
    try:
        application = Application.query.get_or_404(app_id)
        if application.user_session_id != session.get('user_session_id'):
            abort(403)

        resume = application.resume
        if not resume or not resume.structured_text or 'full_text' not in resume.structured_text:
             return jsonify({'error': 'Resume text not found or not cached.'}), 404
        
        scraped_jd = ScrapedJD.query.filter_by(
            user_session_id=session.get('user_session_id'),
            company_name=application.company_name
        ).order_by(ScrapedJD.created_date.desc()).first()
        job_title = scraped_jd.job_title if scraped_jd else f"Role at {application.company_name}"
        
        data = request.get_json() or {}
        custom_prompts = {'interview_prep': data.get('custom_prompt')}
        
        json_structure = """{...}""" # Structure doesn't need to be fully verbose here
        placeholders = {
            'COMPANY': application.company_name,
            'JOB_TITLE': job_title,
            'JOB_DESCRIPTION': application.job_description,
            'FULL_RESUME_TEXT': resume.structured_text['full_text'],
            'JSON_STRUCTURE': json_structure
        }
        prompt_text = processor._get_prompt('interview_prep', custom_prompts, placeholders)
        
        return jsonify({'prompt_text': prompt_text})
    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/scraped-jds', methods=['POST'])
def add_scraped_jd():
    data = request.get_json()
    user_session_id = data.get('user_session_id')
    if not user_session_id:
        return jsonify({'error': 'user_session_id is required'}), 400

    new_jd = ScrapedJD(
        job_title=data.get('job_title', 'N/A'),
        company_name=data.get('company_name', 'N/A'),
        job_description=data.get('job_description', ''),
        page_url=data.get('page_url', ''),
        application_type=data.get('application_type', 'Normal'),  # New field
        user_session_id=user_session_id
    )
    db.session.add(new_jd)
    db.session.commit()
    return jsonify(new_jd.to_dict()), 201

@app.route('/api/scraped-jds', methods=['GET'])
def get_scraped_jds():
    user_session_id = session.get('user_session_id')
    if not user_session_id:
        return jsonify([])
    
    jds = ScrapedJD.query.filter_by(user_session_id=user_session_id).order_by(ScrapedJD.created_date.desc()).all()
    return jsonify([jd.to_dict() for jd in jds])

@app.route('/api/scraped-jds/<int:jd_id>', methods=['DELETE'])
def delete_scraped_jd(jd_id):
    jd = ScrapedJD.query.get_or_404(jd_id)
    if jd.user_session_id != session.get('user_session_id'):
        abort(403)
    db.session.delete(jd)
    db.session.commit()
    return jsonify({'message': 'Scraped JD deleted'})

@app.route('/api/scraped-jds/<int:jd_id>/status', methods=['PUT'])
def update_scraped_jd_status(jd_id):
    jd = ScrapedJD.query.get_or_404(jd_id)
    if jd.user_session_id != session.get('user_session_id'):
        abort(403)
    data = request.get_json()
    if 'status' in data:
        jd.status = data['status']
        db.session.commit()
    return jsonify(jd.to_dict())

@app.route('/api/download_resume', methods=['POST'])
def download_resume_file():
    data = request.get_json()
    data['session_id'] = session['user_session_id']

    # Handle format preference - don't override if already specified
    if 'format' not in data:
        data['format'] = 'pdf'  # Default to PDF only if not specified

    print(f"DEBUG: Download format requested: {data.get('format')}")
    print(f"DEBUG: Full download data: {data}")

    task = celery.send_task('celery_worker.create_download_file_task', args=[data])
    return jsonify({'job_id': task.id})

@app.route('/api/download_cover_letter', methods=['POST'])
def download_cover_letter():
    try:
        data = request.get_json()
        cover_letter_text = data.get('cover_letter_text')
        company_name = data.get('company_name', 'Unknown Company')
        user_name = f"{session.get('user_first_name', '')} {session.get('user_last_name', '')}".strip()

        if not cover_letter_text:
            return jsonify({'error': 'Cover letter text is required'}), 400

        if not user_name:
            user_name = 'Your Name'

        # Create the DOCX file
        docx_path = processor.create_cover_letter_docx(cover_letter_text, company_name, user_name)

        # Generate filename
        safe_company_name = company_name.replace(' ', '_').replace('/', '_')
        filename = f"{user_name.replace(' ', '_')}_{safe_company_name}_Cover_Letter.docx"

        # Move to uploads folder with proper name
        final_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        shutil.move(docx_path, final_path)

        return jsonify({
            'success': True,
            'download_url': f'/download/{filename}',
            'filename': filename
        })

    except Exception as e:
        print(f"Error creating cover letter DOCX: {e}")
        return jsonify({'error': f'Failed to create cover letter document: {str(e)}'}), 500

@app.route('/api/job-status/<job_id>', methods=['GET'])
def get_job_status(job_id):
    """Check the status of a Celery job and return results if completed"""
    try:
        from celery.result import AsyncResult
        task = AsyncResult(job_id, app=celery)

        if task.state == 'PENDING':
            return jsonify({
                'status': 'pending',
                'job_id': job_id,
                'message': 'Job is still processing'
            })
        elif task.state == 'SUCCESS':
            result = task.result
            return jsonify({
                'status': 'completed',
                'job_id': job_id,
                'result': result,
                'message': 'Job completed successfully'
            })
        elif task.state == 'FAILURE':
            return jsonify({
                'status': 'failed',
                'job_id': job_id,
                'error': str(task.info),
                'message': 'Job failed'
            })
        else:
            return jsonify({
                'status': task.state,
                'job_id': job_id,
                'message': f'Job is in {task.state} state'
            })
    except Exception as e:
        return jsonify({
            'status': 'error',
            'job_id': job_id,
            'error': str(e),
            'message': 'Error checking job status'
        })

@app.route('/download/<filename>')
def download_file(filename):
    return send_file(os.path.join(app.config['UPLOAD_FOLDER'], filename), as_attachment=True)

if __name__ == '__main__':
    with app.app_context():
        db.create_all()
    socketio.run(app, debug=True, host='127.0.0.1', port=5001, use_reloader=False)
